"""Generate lab exam Word report with Parts A/B/C, code snippets, screenshot placeholders."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "LAB_EXAM_REPORT.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("\n\n\nBLOCKCHAIN-BASED\nACADEMIC CREDENTIAL VERIFICATION DApp\n")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Lab Examination Report\nDecentralized Credential Verification Platform")
    s.font.size = Pt(14)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("\n\n\nStudent Name: _________________________________\n\n")
    info.add_run("Roll Number: _________________________________\n\n")
    info.add_run("Course / Subject: _________________________________\n\n")
    info.add_run("Date: _________________________________\n\n")
    info.add_run("Instructor: _________________________________\n")
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_code(doc: Document, code: str, caption: str = "") -> None:
    if caption:
        cap = doc.add_paragraph()
        cap.add_run(caption).bold = True
        cap.paragraph_format.space_before = Pt(6)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_screenshot_placeholder(doc: Document, figure_num: str, title: str, instructions: str) -> None:
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.add_run(f"FIGURE {figure_num}: {title}").bold = True
    h.paragraph_format.space_before = Pt(12)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.2)
    set_cell_shading(cell, "F1F5F9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"\n\n[ INSERT SCREENSHOT HERE ]\n\n"
        f"{instructions}\n\n"
        f"(Recommended size: full page width)\n\n"
    )
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.italic = True
    r.font.size = Pt(11)

    cap = doc.add_paragraph(f"Figure {figure_num}: {title}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build_report() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)

    # ── Introduction ──────────────────────────────────────────────────────
    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "This project implements a Blockchain-Based Academic Credential Verification "
        "Decentralized Application (DApp) for a national university consortium. "
        "Universities issue tamper-proof degrees on Ethereum, store certificate documents "
        "on IPFS (Pinata), detect forged or suspicious verification attempts using AI, "
        "log every action in a SHA-256 chained ledger backed by a Merkle Tree, and "
        "display all data on a live web verification dashboard."
    )
    add_body(doc, "End-to-End Workflow:")
    add_bullet(doc, "Wallet Login → Session Verified → Smart Contract Deploy")
    add_bullet(doc, "Credential Registered → Certificate Document on IPFS (Pinata)")
    add_bullet(doc, "AI Forgery/Threat Check → Transaction Hashed & Chained")
    add_bullet(doc, "Merkle Root Generated → Web Dashboard Rendered")

    add_heading(doc, "1.1 Technology Stack", 2)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Blockchain Platform", "Ethereum Sepolia Testnet"),
        ("Smart Contract", "Solidity ^0.8.20 (CredentialRegistry.sol)"),
        ("Chain Interaction", "Web3.py"),
        ("Backend", "Python 3.11"),
        ("Frontend", "Flask + HTML/CSS/JS Web Dashboard"),
        ("Storage", "Pinata IPFS + SHA-256 Content Identifier (CID)"),
        ("Cryptography", "ECDSA key pairs + SHA-256 hashing"),
        ("AI / ML", "scikit-learn IsolationForest + rule-based detection"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    add_heading(doc, "1.2 How to Run the Project", 2)
    add_code(
        doc,
        "cd c:\\BLOCKCHAINFINAL\n"
        "pip install -r requirements.txt\n"
        "python scripts\\compile_contract.py\n"
        "python server.py          REM Opens web dashboard at http://127.0.0.1:5000",
        "Listing 1: Commands to run the full DApp",
    )

    add_screenshot_placeholder(
        doc, "1", "Project Folder in VS Code",
        "Screenshot of VS Code with BLOCKCHAINFINAL folder open showing "
        "contracts/, backend/, web/ directories.",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # PART A
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART A — Wallet-Based Identity System & Smart Contract Registry (15 Marks)", 1)

    add_heading(doc, "A.1 Objectives", 2)
    add_bullet(doc, "Generate ECDSA key pairs for Registrar, Student, and Employer roles")
    add_bullet(doc, "Implement wallet login with session tokens (VERIFIED / FAILED / EXPIRED)")
    add_bullet(doc, "Deploy CredentialRegistry.sol on Sepolia and issue/revoke credentials")
    add_bullet(doc, "Enforce onlyRegistrar modifier inside Solidity contract")

    add_heading(doc, "A.2 Deterministic ECDSA Identity Generation", 2)
    add_body(
        doc,
        "Three wallet identities are generated at runtime using deterministic seeds. "
        "Each identity contains role, wallet_address (first 20 bytes of SHA-256 public key hash), "
        "public_key, private_key, session_token, and token_expiry."
    )
    add_code(
        doc,
        '''def _derive_signing_key(seed_string: str) -> SigningKey:
    """Deterministically derive SECP256k1 key from fixed seed string."""
    digest = hashlib.sha256(seed_string.encode("utf-8")).digest()
    return SigningKey.from_string(digest, curve=SECP256k1)

def _pubkey_to_wallet(public_key_bytes: bytes) -> str:
    pubkey_hash = hashlib.sha256(public_key_bytes).digest()
    return "0x" + pubkey_hash[:20].hex()''',
        "Listing 2: Deterministic key derivation (backend/part_a/identity.py)",
    )

    add_screenshot_placeholder(
        doc, "2", "Identity Table & Login Results",
        "Terminal or dashboard Panel 1 showing:\n"
        "• Identity table (Role, Wallet Address, Public Key)\n"
        "• Login results: VERIFIED / FAILED / EXPIRED\n"
        "• Active session token (first 16 chars)\n"
        "• Role-permission mapping table",
    )

    add_heading(doc, "A.3 Login and Session Verification", 2)
    add_body(
        doc,
        "The login() function accepts wallet address and private key, validates credentials, "
        "and returns a signed session token. verify_session() is called before every "
        "subsequent operation in Parts B and C."
    )
    add_code(
        doc,
        '''def login(self, wallet_address: str, private_key: str) -> LoginResult:
    ident = self._identities_by_wallet.get(wallet_address.lower())
    if ident is None:
        return LoginResult(False, "FAILED", "Wallet address not registered")
    if private_key != ident.private_key:
        return LoginResult(False, "FAILED", "Invalid private key")
    session_token = hashlib.sha256(f"{ident.role}:{time.time()}".encode()).hexdigest()
    return LoginResult(True, "VERIFIED", "Session established", session_token=session_token)

def verify_session(self, token: str):
    session = self.sessions.get(token)
    if session is None:
        return False, "Invalid session token", None
    if time.time() > session.token_expiry:
        return False, "Session token expired", None
    return True, "Session valid", session''',
        "Listing 3: Login and verify_session (backend/part_a/auth.py)",
    )

    add_heading(doc, "A.4 Smart Contract — CredentialRegistry.sol", 2)
    add_body(
        doc,
        "The Solidity smart contract stores credentials on Ethereum with four main functions: "
        "issueCredential, revokeCredential, transferCredentialOwnership, and getCredential. "
        "Access control is enforced via onlyRegistrar modifier inside the contract."
    )
    add_code(
        doc,
        '''modifier onlyRegistrar() {
    require(isRegistrar[msg.sender], "Only registrar can perform this action");
    _;
}

function issueCredential(
    string calldata studentName,
    address registrar,
    string calldata documentCid
) external onlyRegistrar returns (uint256) {
    credentialCount++;
    credentials[credentialCount] = Credential({... status: Status.VALID ...});
    emit CredentialIssued(credentialCount, studentName, registrar, documentCid, block.timestamp);
    return credentialCount;
}

function verifyCredential(uint256 credentialId) external view returns (...) {
    require(cred.status == Status.VALID, "Credential is REVOKED and cannot be verified");
}''',
        "Listing 4: Smart contract core functions (contracts/CredentialRegistry.sol)",
    )

    add_screenshot_placeholder(
        doc, "3", "Sepolia Contract Deployment on Etherscan",
        "Screenshot of Sepolia Etherscan showing:\n"
        "• Contract address (0x...)\n"
        "• Deployment transaction hash\n"
        "• Contract bytecode present (not empty)\n"
        "URL: https://sepolia.etherscan.io/address/YOUR_CONTRACT_ADDRESS",
    )

    add_heading(doc, "A.5 Required Output — Part A", 2)
    add_bullet(doc, "2 credential issuance records (Alice Johnson — VALID, Bob Smith — issued then REVOKED)")
    add_bullet(doc, "1 successful revoke result for credential ID 2")
    add_bullet(doc, "1 revert message when verifying revoked credential")

    add_screenshot_placeholder(
        doc, "4", "Issued Credentials & Revoke Demonstration",
        "Terminal output OR Dashboard Panel 2 showing:\n"
        "• Credential ID 1 — Alice Johnson — VALID\n"
        "• Credential ID 2 — Bob Smith — REVOKED\n"
        "• Revoke transaction hash\n"
        "• Revert message: 'Credential is REVOKED and cannot be verified'",
    )

    add_screenshot_placeholder(
        doc, "5", "Remix IDE — Contract Deployment (Optional)",
        "Screenshot of Remix IDE showing:\n"
        "• CredentialRegistry.sol compiled\n"
        "• Deployed via Injected Provider (MetaMask)\n"
        "• Contract functions visible in Remix",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # PART B
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART B — IPFS Document Storage, Integrity & AI Forgery Detection (15 Marks)", 1)

    add_heading(doc, "B.1 IPFS Document Storage (Pinata)", 2)
    add_body(
        doc,
        "Certificate documents are JSON files describing degrees. SHA-256 digest of each "
        "document is computed and used as the Content Identifier (CID). Documents are "
        "uploaded to Pinata IPFS and the CID is passed to issueCredential() on the smart contract."
    )
    add_code(
        doc,
        '''@staticmethod
def compute_cid(content: dict[str, Any]) -> str:
    raw = json.dumps(content, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()

def verify_document(self, cid: str) -> tuple[str, str]:
    stored = self.store[cid]
    recomputed = self.compute_cid(stored)
    if recomputed == cid:
        return "INTACT", "Document hash matches CID"
    return "TAMPERED", f"Hash mismatch for CID {cid[:20]}..."''',
        "Listing 5: IPFS CID computation and integrity check (backend/part_b/ipfs_storage.py)",
    )

    add_screenshot_placeholder(
        doc, "6", "Document-to-CID Mapping & Pinata Upload",
        "Terminal output OR dashboard IPFS section showing:\n"
        "• 3 certificate files: alice_bsc.json, bob_msc.json, carol_phd_spare.json\n"
        "• SHA-256 CID for each (computed, not hardcoded)\n"
        "• Pinata IPFS hash for each document",
    )

    add_screenshot_placeholder(
        doc, "7", "Pinata Cloud — Pinned Files",
        "Screenshot of Pinata dashboard (pinata.cloud) showing:\n"
        "• 3 pinned JSON certificate files\n"
        "• IPFS hashes visible\n"
        "• Gateway links to view documents",
    )

    add_heading(doc, "B.2 Tampering Demonstration (Mandatory)", 2)
    add_body(
        doc,
        "One byte/field of carol_phd_spare.json is modified. verify_document() is called "
        "and returns TAMPERED with the affected CID identified. A security event is "
        "forwarded to the AI threat detection layer."
    )
    add_screenshot_placeholder(
        doc, "8", "Tampered Document Detection",
        "Terminal output showing:\n"
        "• Tampered document CID clearly identified\n"
        "• verify_document() result: TAMPERED\n"
        "• Reason string: Hash mismatch for CID...\n"
        "• Integrity table: carol_phd_spare.json = TAMPERED, others = INTACT",
    )

    add_heading(doc, "B.3 AI Threat Detection — IsolationForest + Rules", 2)
    add_body(
        doc,
        "Security events (failed logins, expired sessions, revoked-credential reverts, "
        "tampered documents) are collected. An IsolationForest model is trained at runtime "
        "on 20+ records. Rule-based layer flags HIGH if failed_logins > 2 or doc_tamper_flag == 1; "
        "MEDIUM if verification_attempts > 3."
    )
    add_code(
        doc,
        '''def _rule_based_level(self, row: dict) -> str:
    if row["failed_logins"] > 2 or row["doc_tamper_flag"] == 1:
        return "HIGH"
    if row["verification_attempts"] > 3:
        return "MEDIUM"
    return "LOW"

model = IsolationForest(contamination=0.15, random_state=42)
model.fit(X_train)
final_level = max(ml_level, rule_level)  # higher of ML and rule result''',
        "Listing 6: AI threat detection (backend/part_b/threat_detection.py)",
    )

    add_screenshot_placeholder(
        doc, "9", "AI Threat Detection Results",
        "Terminal output OR Dashboard Panel 3 showing:\n"
        "• Training dataset size (≥ 20 records)\n"
        "• Model accuracy (%)\n"
        "• Per-transaction threat table: tx_id, threat level, confidence score\n"
        "• List of HIGH alerts (including forgery/tamper alerts)",
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # PART C
    # ══════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART C — SHA-256 Chained Ledger, Merkle Tree & Live Dashboard (20 Marks)", 1)

    add_heading(doc, "C.1 SHA-256 Chained Transaction Ledger", 2)
    add_body(
        doc,
        "Every event from Parts A and B is logged as a transaction object with fields: "
        "tx_id, action, actor_wallet, timestamp, payload_hash, prev_hash. "
        "Genesis block uses prev_hash = '0' * 64. Chain includes 10+ real runtime transactions."
    )
    add_code(
        doc,
        '''@dataclass
class Transaction:
    tx_id: str
    action: str
    actor_wallet: str
    payload_hash: str
    prev_hash: str
    data: dict

    def compute_hash(self) -> str:
        payload = json.dumps(self.data, sort_keys=True)
        return hashlib.sha256(payload.encode()).hexdigest()

# Chain verification: each prev_hash must match previous payload_hash''',
        "Listing 7: Chained ledger (backend/part_c/ledger.py)",
    )

    add_screenshot_placeholder(
        doc, "10", "Full Transaction Log & Chain Verification",
        "Terminal output showing:\n"
        "• Transaction log table: tx_id | action | payload_hash (16) | prev_hash (16)\n"
        "• Chain verification: VALID per link (32+ transactions)\n"
        "• Total transactions ≥ 10",
    )

    add_heading(doc, "C.2 Merkle Tree (Built from Scratch)", 2)
    add_body(
        doc,
        "A Merkle Tree is built from all payload_hash values without using any external "
        "tree library. Transaction #3 is tampered to show root change. Merkle Proof for "
        "transaction #1 is generated and verified."
    )
    add_code(
        doc,
        '''def _hash_pair(left: str, right: str) -> str:
    return hashlib.sha256((left + right).encode()).hexdigest()

class MerkleTree:
    def _build(self) -> str:
        layer = list(self.leaves)
        while len(layer) > 1:
            next_layer = []
            for i in range(0, len(layer), 2):
                left = layer[i]
                right = layer[i + 1] if i + 1 < len(layer) else layer[i]
                next_layer.append(_hash_pair(left, right))
            layer = next_layer
        return layer[0]''',
        "Listing 8: Merkle tree implementation (backend/part_c/merkle.py)",
    )

    add_screenshot_placeholder(
        doc, "11", "Merkle Root, Tamper Demo & Proof Verification",
        "Terminal output showing:\n"
        "• Original Merkle root (full hash)\n"
        "• Tampered Merkle root (different from original)\n"
        "• Changed-node path from tampered leaf to root\n"
        "• Merkle Proof verification result: VALID",
    )

    add_heading(doc, "C.3 Live Web Verification Dashboard", 2)
    add_body(
        doc,
        "A professional web dashboard at http://127.0.0.1:5000 displays 4 live panels "
        "with data from the Python backend API (no hardcoded values). The Verify Credential "
        "button allows an employer to enter a credential ID and receive VALID / REVOKED / "
        "NOT FOUND results with real-time panel refresh."
    )

    add_screenshot_placeholder(
        doc, "12", "Web Dashboard — Panel 1 & 2 (Identities + Credentials)",
        "Full browser screenshot of http://127.0.0.1:5000 showing:\n"
        "• Panel 1: Wallet Identities (role, wallet address, session status)\n"
        "• Panel 2: Issued Credentials (ID, student, registrar, VALID/REVOKED)\n"
        "• Stats row: Contract address, Merkle root, chain integrity, AI accuracy",
    )

    add_screenshot_placeholder(
        doc, "13", "Web Dashboard — Panel 3 & 4 (Threats + Transaction Chain)",
        "Browser screenshot showing:\n"
        "• Panel 3: Threat Alert Feed (tx_id, threat badge, confidence, reason)\n"
        "• Panel 4: Transaction Chain & Merkle Root Status",
    )

    add_heading(doc, "C.4 Verify Credential Demonstration", 2)
    add_body(
        doc,
        "Employer role is selected. Credential ID 1 returns VALID. Credential ID 2 returns "
        "REVOKED. Dashboard panels refresh after each verification."
    )

    add_screenshot_placeholder(
        doc, "14", "Verify Credential — VALID Result (ID = 1)",
        "Dashboard screenshot showing:\n"
        "• Role: Employer (Verifier)\n"
        "• Credential ID: 1\n"
        "• Result: VALID (green badge)\n"
        "• Toast notification: 'Credential #1: VALID'",
    )

    add_screenshot_placeholder(
        doc, "15", "Verify Credential — REVOKED Result (ID = 2)",
        "Dashboard screenshot showing:\n"
        "• Role: Employer (Verifier)\n"
        "• Credential ID: 2\n"
        "• Result: REVOKED (red badge)\n"
        "• Details: Credential is REVOKED and cannot be verified",
    )

    add_screenshot_placeholder(
        doc, "16", "MetaMask — Sepolia Transactions (Optional)",
        "Screenshot of MetaMask showing:\n"
        "• Sepolia testnet selected\n"
        "• Account balance\n"
        "• Recent transactions (deploy, issue, revoke)",
    )

    doc.add_page_break()

    # ── Conclusion ────────────────────────────────────────────────────────
    add_heading(doc, "4. Conclusion", 1)
    add_body(
        doc,
        "This DApp successfully demonstrates a complete academic credential verification "
        "pipeline: wallet-based authentication, Ethereum smart contract registry on Sepolia, "
        "IPFS document storage via Pinata with SHA-256 integrity verification, AI-powered "
        "threat detection using IsolationForest and rule-based classification, SHA-256 chained "
        "audit ledger with Merkle tree integrity proofs, and a live professional web dashboard "
        "for real-time credential verification by employers worldwide."
    )

    add_heading(doc, "5. Screenshot Checklist", 2)
    checklist = doc.add_table(rows=17, cols=3)
    checklist.style = "Table Grid"
    hdr = checklist.rows[0].cells
    hdr[0].text = "Fig #"
    hdr[1].text = "What to Screenshot"
    hdr[2].text = "Done ✓"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True

    items = [
        ("1", "VS Code project folder"),
        ("2", "Identity table & login results"),
        ("3", "Sepolia Etherscan contract page"),
        ("4", "Issued credentials & revoke output"),
        ("5", "Remix IDE deployment (optional)"),
        ("6", "Document-to-CID mapping"),
        ("7", "Pinata pinned files"),
        ("8", "Tampered document TAMPERED result"),
        ("9", "AI threat table & HIGH alerts"),
        ("10", "Transaction log & chain VALID"),
        ("11", "Merkle root, tamper path, proof VALID"),
        ("12", "Dashboard Panels 1 & 2"),
        ("13", "Dashboard Panels 3 & 4"),
        ("14", "Verify Credential ID=1 VALID"),
        ("15", "Verify Credential ID=2 REVOKED"),
        ("16", "MetaMask Sepolia (optional)"),
    ]
    for i, (num, desc) in enumerate(items, 1):
        checklist.rows[i].cells[0].text = num
        checklist.rows[i].cells[1].text = desc
        checklist.rows[i].cells[2].text = "☐"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Report saved: {OUT}")


if __name__ == "__main__":
    build_report()
